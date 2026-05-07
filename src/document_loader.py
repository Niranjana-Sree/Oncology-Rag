"""
Document loader for MedRAG.
Supports PDF, DOCX, TXT, and HTML formats.
Cleans headers, footers, page numbers, and other noise from extracted text.
"""

from __future__ import annotations

import hashlib
import logging
import re
import unicodedata
from pathlib import Path
from typing import Optional

logger = logging.getLogger("medrag.document_loader")


# ---------------------------------------------------------------------------
# Format-specific extractors
# ---------------------------------------------------------------------------

def _extract_pdf(file_path: Path) -> tuple[str, int]:
    """Extract raw text from a PDF file. Returns (text, page_count)."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF loading. Run: pip install PyPDF2")

    pages: list[str] = []
    with open(file_path, "rb") as fh:
        reader = PyPDF2.PdfReader(fh)
        page_count = len(reader.pages)
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

    logger.debug("PDF '%s' — %d pages extracted", file_path.name, page_count)
    return "\n".join(pages), page_count


def _extract_docx(file_path: Path) -> tuple[str, int]:
    """Extract raw text from a DOCX file. Returns (text, paragraph_count)."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required for DOCX loading. Run: pip install python-docx")

    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    logger.debug("DOCX '%s' — %d paragraphs extracted", file_path.name, len(paragraphs))
    return text, len(paragraphs)


def _extract_txt(file_path: Path) -> tuple[str, int]:
    """Read a plain-text file. Returns (text, line_count)."""
    text = file_path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    logger.debug("TXT '%s' — %d lines read", file_path.name, len(lines))
    return text, len(lines)


def _extract_html(file_path: Path) -> tuple[str, int]:
    """Strip HTML tags and return clean text. Returns (text, tag_count)."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("beautifulsoup4 is required for HTML loading. Run: pip install beautifulsoup4")

    raw_html = file_path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw_html, "lxml")

    # Remove script, style, nav, header, footer elements
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    tag_count = len(soup.find_all())
    logger.debug("HTML '%s' — tags removed, text extracted", file_path.name)
    return text, tag_count


# ---------------------------------------------------------------------------
# Noise cleaning
# ---------------------------------------------------------------------------

# Patterns that mark lines as noise
_PAGE_NUMBER_RE = re.compile(
    r"^\s*"
    r"(?:"
    r"(?:page\s*)?[-–]?\s*\d{1,4}\s*[-–]?"    # "Page 3", "- 3 -", "3"
    r"|"
    r"page\s+\d{1,4}\s+of\s+\d{1,4}"           # "Page 3 of 120"
    r")"
    r"\s*$",
    re.IGNORECASE,
)

_SHORT_REPEATED_RE = re.compile(r"^(.{1,60})\1{2,}$")   # same short phrase repeated 3+ times

_RUNNING_HEADER_FOOTER_RE = re.compile(
    r"^\s*(?:"
    r"confidential|draft|proprietary|internal use only"
    r"|all rights reserved|copyright\s*©?"
    r"|www\.[^\s]{3,}|https?://"
    r")\s*$",
    re.IGNORECASE,
)

_WHITESPACE_LINE_RE = re.compile(r"^\s*$")


def _is_noise_line(line: str) -> bool:
    """Return True if a line is a header/footer/page-number artefact."""
    stripped = line.strip()
    if not stripped:
        return True
    if _PAGE_NUMBER_RE.match(stripped):
        return True
    if _RUNNING_HEADER_FOOTER_RE.match(stripped):
        return True
    if _SHORT_REPEATED_RE.match(stripped):
        return True
    # Very short lines that are likely artefacts (< 4 chars, not a list bullet)
    if len(stripped) < 4 and not stripped.startswith(("-", "•", "*", "·")):
        return True
    return False


def _clean_text(raw: str) -> str:
    """
    Clean extracted document text:
    1. Normalise unicode to NFC
    2. Replace non-breaking spaces and other invisible characters
    3. Drop noise lines (headers, footers, page numbers)
    4. Collapse excessive blank lines (max 2 consecutive)
    5. Strip leading/trailing whitespace from each line
    """
    # 1. Unicode normalisation
    text = unicodedata.normalize("NFC", raw)

    # 2. Replace common invisible / non-printing characters
    text = text.replace(" ", " ")   # non-breaking space
    text = text.replace("​", "")    # zero-width space
    text = text.replace("�", "")    # replacement character
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)  # control chars

    # 3. Drop noise lines
    lines = text.splitlines()
    clean_lines: list[str] = []
    for line in lines:
        if not _is_noise_line(line):
            clean_lines.append(line.rstrip())

    # 4. Collapse 3+ consecutive blank lines down to 2
    collapsed: list[str] = []
    blank_count = 0
    for line in clean_lines:
        if _WHITESPACE_LINE_RE.match(line):
            blank_count += 1
            if blank_count <= 2:
                collapsed.append("")
        else:
            blank_count = 0
            collapsed.append(line)

    # 5. Final strip
    return "\n".join(collapsed).strip()


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".txt", ".html", ".htm"}


def _make_doc_id(file_path: Path) -> str:
    """Create a stable doc_id from the file path using a short hash."""
    digest = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
    stem = re.sub(r"[^a-zA-Z0-9_-]", "_", file_path.stem)[:40]
    return f"{stem}__{digest}"


def load_document(
    file_path: str | Path,
    category: str,
    extra_metadata: Optional[dict] = None,
) -> "RawDocument":
    """
    Load and clean a single document from disk.

    Args:
        file_path:  Path to the document file.
        category:   DocCategory value (e.g. "drug_monograph").
        extra_metadata: Any additional key-value pairs to store on the document.

    Returns:
        A RawDocument instance with cleaned text and metadata.

    Raises:
        FileNotFoundError: if the file does not exist.
        ValueError: if the file extension is not supported.
    """
    # Import here to avoid circular imports at module level
    from src.models import RawDocument

    path = Path(file_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info("Loading %s — %s", ext.upper(), path.name)

    # Dispatch to format-specific extractor
    if ext == ".pdf":
        raw_text, page_count = _extract_pdf(path)
    elif ext == ".docx":
        raw_text, page_count = _extract_docx(path)
    elif ext in {".txt"}:
        raw_text, page_count = _extract_txt(path)
    elif ext in {".html", ".htm"}:
        raw_text, page_count = _extract_html(path)
    else:
        raw_text, page_count = _extract_txt(path)

    cleaned = _clean_text(raw_text)
    word_count = len(cleaned.split())

    metadata: dict = {
        "file_name": path.name,
        "file_ext": ext,
        "category": category,
        **(extra_metadata or {}),
    }

    doc = RawDocument(
        doc_id=_make_doc_id(path),
        file_path=str(path),
        category=category,
        raw_text=cleaned,
        metadata=metadata,
        page_count=page_count,
        word_count=word_count,
    )

    logger.info(
        "Loaded '%s' → doc_id=%s | words=%d | pages/paras=%d",
        path.name, doc.doc_id, word_count, page_count,
    )
    return doc


def load_directory(
    directory: str | Path,
    category: str,
    recursive: bool = True,
) -> list["RawDocument"]:
    """
    Load all supported documents from a directory.

    Args:
        directory:  Path to the directory to scan.
        category:   DocCategory value applied to all documents in this directory.
        recursive:  If True, descend into subdirectories.

    Returns:
        List of RawDocument instances, one per file.
    """
    dir_path = Path(directory).resolve()
    if not dir_path.exists():
        logger.warning("Directory does not exist: %s", dir_path)
        return []

    pattern = "**/*" if recursive else "*"
    files = [
        f for f in dir_path.glob(pattern)
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    if not files:
        logger.warning("No supported documents found in: %s", dir_path)
        return []

    logger.info("Found %d documents in '%s'", len(files), dir_path.name)

    documents: list = []
    for f in sorted(files):
        try:
            doc = load_document(f, category=category)
            documents.append(doc)
        except Exception as exc:
            logger.error("Failed to load '%s': %s", f.name, exc)

    logger.info("Successfully loaded %d / %d documents from '%s'",
                len(documents), len(files), dir_path.name)
    return documents


def load_all_data_dirs() -> list["RawDocument"]:
    """
    Load all documents from the oncology corpus data subdirectories.
    Categories: clinical_guideline, research_paper, general.
    (drug_monograph removed — oncology-only corpus)
    """
    import config
    from src.models import DocCategory

    all_docs: list = []
    mapping = [
        (config.DATA_CLINICAL_GUIDELINES, DocCategory.CLINICAL_GUIDELINE.value),
        (config.DATA_RESEARCH_PAPERS,     DocCategory.RESEARCH_PAPER.value),
        (config.DATA_GENERAL,             DocCategory.GENERAL.value),
    ]

    for data_dir, category in mapping:
        docs = load_directory(data_dir, category=category)
        all_docs.extend(docs)

    logger.info("Total documents loaded across all categories: %d", len(all_docs))
    return all_docs
